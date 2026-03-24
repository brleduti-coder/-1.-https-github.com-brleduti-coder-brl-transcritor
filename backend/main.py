import asyncio
import os
import re
import shutil
import time
from pathlib import Path

import sys

import httpx
from dotenv import load_dotenv
from fastapi import BackgroundTasks, FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

load_dotenv()

ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY", "")
FRONTEND_ORIGIN = os.getenv("FRONTEND_ORIGIN", "http://localhost:3000")
BASE_URL = os.getenv("BASE_URL", "http://localhost:8000")

VENV_BIN = Path(sys.executable).parent
YT_DLP = str(VENV_BIN / "yt-dlp")
DOWNLOAD_DIR = Path(__file__).parent / "downloads"
TRANSCRICOES_DIR = Path(__file__).parent / "transcricoes"
DOWNLOAD_DIR.mkdir(exist_ok=True)
TRANSCRICOES_DIR.mkdir(exist_ok=True)

app = FastAPI(title="BRL Transcritor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[FRONTEND_ORIGIN, "http://localhost:5173"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory job tracking
jobs: dict[str, dict] = {}


class TranscriptionRequest(BaseModel):
    url: str
    title: str = "Título não identificado"


def _sanitize_filename(name: str) -> str:
    """Remove characters that are unsafe for filenames."""
    name = re.sub(r'[\\/:*?"<>|]', "", name)
    name = name.strip(". ")
    return name or "transcricao"


# ---------------------------------------------------------------------------
# Core pipeline
# ---------------------------------------------------------------------------

async def _process_video(job_id: str, url: str, title: str):
    """Full pipeline: download → upload → transcribe → save .txt → cleanup video."""
    job = jobs[job_id]
    video_path: Path | None = None

    try:
        # 1. Download video with yt-dlp
        job["status"] = "downloading"
        uid = str(int(time.time() * 1000))
        output_template = str(DOWNLOAD_DIR / f"{uid}_%(title)s.%(ext)s")

        proc = await asyncio.create_subprocess_exec(
            YT_DLP,
            "--no-playlist",
            "--js-runtimes", "nodejs",
            "-o", output_template,
            url,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await proc.communicate()

        if proc.returncode != 0:
            raise RuntimeError(f"yt-dlp falhou: {stderr.decode()}")

        # Find downloaded file
        matches = sorted(DOWNLOAD_DIR.glob(f"{uid}_*"), key=lambda p: p.stat().st_mtime)
        if not matches:
            raise RuntimeError("Arquivo de vídeo não encontrado após download.")
        video_path = matches[-1]

        # Extract video title from filename (uid_TITLE.ext)
        video_title = video_path.stem  # e.g. "1234_My Video Title"
        if video_title.startswith(f"{uid}_"):
            video_title = video_title[len(f"{uid}_"):]
        video_title = video_title or title

        # 2. Upload to AssemblyAI
        job["status"] = "uploading"
        async with httpx.AsyncClient(timeout=300) as client:
            with open(video_path, "rb") as f:
                upload_resp = await client.post(
                    "https://api.assemblyai.com/v2/upload",
                    headers={
                        "Authorization": ASSEMBLYAI_API_KEY,
                        "Content-Type": "application/octet-stream",
                    },
                    content=f.read(),
                )
            upload_resp.raise_for_status()
            upload_url = upload_resp.json()["upload_url"]

        # 3. Start transcription
        job["status"] = "transcribing"
        async with httpx.AsyncClient(timeout=60) as client:
            transcript_resp = await client.post(
                "https://api.assemblyai.com/v2/transcript",
                headers={
                    "Authorization": ASSEMBLYAI_API_KEY,
                    "Content-Type": "application/json",
                },
                json={
                    "audio_url": upload_url,
                    "language_code": "pt",
                    "speech_models": ["universal-2"],
                },
            )
            transcript_resp.raise_for_status()
            transcript_id = transcript_resp.json()["id"]

        # 4. Poll until completed
        async with httpx.AsyncClient(timeout=30) as client:
            while True:
                await asyncio.sleep(15)
                status_resp = await client.get(
                    f"https://api.assemblyai.com/v2/transcript/{transcript_id}",
                    headers={"Authorization": ASSEMBLYAI_API_KEY},
                )
                status_resp.raise_for_status()
                data = status_resp.json()

                if data["status"] == "completed":
                    transcription_text = data["text"]
                    break
                elif data["status"] == "error":
                    raise RuntimeError(
                        f"Transcrição falhou: {data.get('error', 'erro desconhecido')}"
                    )

        # 5. Save transcription as .docx file
        job["status"] = "saving"
        from docx import Document

        safe_title = _sanitize_filename(video_title)
        docx_filename = f"{safe_title}.docx"
        docx_path = TRANSCRICOES_DIR / docx_filename

        doc = Document()
        doc.add_heading(video_title, level=1)
        doc.add_paragraph(transcription_text)
        doc.save(str(docx_path))

        download_url = f"{BASE_URL}/download/{docx_filename}"

        job.update({
            "status": "completed",
            "message": "Transcrição feita com sucesso!",
            "download_url": download_url,
            "filename": docx_filename,
        })

    except Exception as exc:
        job.update({
            "status": "error",
            "message": str(exc),
        })

    finally:
        # 6. Cleanup video file
        if video_path and video_path.exists():
            video_path.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# File server
# ---------------------------------------------------------------------------

@app.get("/download/{filename}")
async def download_file(filename: str):
    """Serve a transcription file for download."""
    file_path = TRANSCRICOES_DIR / filename
    if not file_path.exists() or not file_path.is_relative_to(TRANSCRICOES_DIR):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado.")
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/transcricoes")
async def list_transcricoes():
    """List all available transcription files."""
    files = []
    for f in sorted(TRANSCRICOES_DIR.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True):
        files.append({
            "filename": f.name,
            "download_url": f"{BASE_URL}/download/{f.name}",
            "size_bytes": f.stat().st_size,
            "created_at": f.stat().st_mtime,
        })
    return {"files": files}


# ---------------------------------------------------------------------------
# Transcription endpoints
# ---------------------------------------------------------------------------

@app.post("/webhook/transcrever-video")
async def transcrever_video(req: TranscriptionRequest, background_tasks: BackgroundTasks):
    """Receives a video URL, processes in background, returns a job ID."""
    if not ASSEMBLYAI_API_KEY:
        raise HTTPException(status_code=500, detail="ASSEMBLYAI_API_KEY não configurada.")
    if not shutil.which(YT_DLP):
        raise HTTPException(status_code=500, detail="yt-dlp não encontrado no sistema.")

    job_id = str(int(time.time() * 1000))
    jobs[job_id] = {"status": "queued", "message": "", "download_url": "", "filename": ""}

    background_tasks.add_task(_process_video, job_id, req.url, req.title)
    return {"job_id": job_id, "status": "queued"}


@app.post("/webhook/transcrever-video/sync")
async def transcrever_video_sync(req: TranscriptionRequest):
    """Synchronous version — waits for the full pipeline before responding."""
    if not ASSEMBLYAI_API_KEY:
        raise HTTPException(status_code=500, detail="ASSEMBLYAI_API_KEY não configurada.")
    if not shutil.which(YT_DLP):
        raise HTTPException(status_code=500, detail="yt-dlp não encontrado no sistema.")

    job_id = str(int(time.time() * 1000))
    jobs[job_id] = {"status": "queued", "message": "", "download_url": "", "filename": ""}

    await _process_video(job_id, req.url, req.title)

    job = jobs[job_id]
    if job["status"] == "error":
        raise HTTPException(status_code=500, detail=job["message"])

    return {
        "message": job["message"],
        "link": job["download_url"],
    }


@app.get("/webhook/transcrever-video/status/{job_id}")
async def get_job_status(job_id: str):
    """Poll job status (for the async endpoint)."""
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job não encontrado.")
    return {"job_id": job_id, **job}


@app.get("/health")
async def health():
    return {"status": "ok", "yt_dlp": shutil.which(YT_DLP) is not None}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
